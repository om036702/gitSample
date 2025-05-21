from docx import Document

# スタイル付きの空のdocxを読み込む（テンプレートのように扱う）
doc = Document("tempdoc.docx")
ctit=["A ","B ","C ","D "]
# スタイル名は .docx 内に定義されているものをそのまま使える
#doc.add_paragraph("【問1】", style="QuestionNumber")
doc.add_paragraph("１. 問題文の内容です。", style="QuestionText")

choices = ["ア. りんご", "イ. みかん", "ウ. ぶどう", "エ. もも"]
for ch in choices:
    doc.add_paragraph(ch, style="chice1")

choi2=["対象の生活空間は、「私的（プライベート）空間」と「公共（パブリック）空間」に大別される",
"生産財は業務用ともいわれ、医療機器など専門家向け商品のことを指す",
"第二次産業の「情報通信業・サービス業」にも関わりは広がっている",
"公共空間のベンチ・街路灯などのストリートファニチャーは「環境デザイン」と呼ばれる"]


doc.add_paragraph("２. 人間生活環境を形成するデザインの対象として、次の中に含まれていないものはどれか選んでください。", style="QuestionText")

for i,ch in enumerate(choi2):
    t1=f"{ctit[i]} {ch}"
    doc.add_paragraph(f"{ctit[i]} {ch}", style="chice1")

doc.save("generated_exam.docx")